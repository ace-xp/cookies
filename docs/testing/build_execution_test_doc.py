from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "策略对接品牌广告与图文广告全链路测试用例-执行版.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "5F6B7A"
LINE = "CBD5E1"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "E3F4EA"
GREEN_TEXT = "226A45"
AMBER = "FFF4D6"
AMBER_TEXT = "7A5A00"
RED = "FDE8E8"
RED_TEXT = "9B1C1C"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, ascii_font="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, ascii_font="Calibri", east_asia="Microsoft YaHei", size=11,
                   color=INK, bold=False):
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)


def style_cell_text(cell, color=INK, size=9.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.5,
              center_columns=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, header_fill)
        style_cell_text(cell, color=NAVY, size=9.3, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = str(value)
            style_cell_text(cell, size=font_size,
                            align=WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, label, text, fill=CALLOUT, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10, color=label_color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def next_numbering_ids(numbering):
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc, kind="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_ids(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    level.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    level.append(ppr)
    if kind == "bullet":
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        rpr.append(fonts)
        level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_list(doc, items, kind="decimal", after=4):
    num_id = create_numbering(doc, kind)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.25
        ppr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        n_id = OxmlElement("w:numId")
        n_id.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(n_id)
        ppr.append(num_pr)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)
    return num_id


def add_label_paragraph(doc, label, text, label_color=BLUE, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.3, color=label_color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.3, color=INK)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("cookies  |  全链路测试执行手册")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = hp.add_run("                                                   执行版 v2.0")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)

    doc.core_properties.title = "策略对接品牌广告与图文广告全链路测试执行手册"
    doc.core_properties.subject = "Strategy → Creative 端到端人工验收"
    doc.core_properties.author = "cookies QA"
    doc.core_properties.comments = "重构为目标、操作、预期结果和证据驱动的执行版。"


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("QUALITY ASSURANCE PLAYBOOK · EXECUTION EDITION")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("策略对接品牌广告与图文广告")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("全链路测试执行手册")
    set_run_font(r, size=17, color=DARK_BLUE, bold=True)

    add_table(
        doc,
        ["版本", "默认演示环境", "更新日期", "执行范围"],
        [["v2.0", "14.103.24.58:8091（HTTP 测试环境）", "2026-08-12", "Strategy → Brand Video / Image Text"]],
        [1200, 3420, 1500, 3240],
        header_fill=LIGHT_BLUE,
        font_size=9.5,
        center_columns=(0, 2),
    )

    add_callout(
        doc,
        "怎么使用",
        "按链路顺序逐阶段执行。每个阶段先看“测试什么”，再照“怎么做”操作；只有“应该看到”的结果全部满足，才进入下一阶段。",
        fill=CALLOUT,
    )

    add_table(
        doc,
        ["2 条", "18 个", "4 类", "1 份"],
        [["核心业务链路", "主链路阶段", "前置门禁", "统一证据台账"]],
        [2340, 2340, 2340, 2340],
        header_fill=AMBER,
        font_size=9.5,
        center_columns=(0, 1, 2, 3),
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("品牌广告主案例")
    set_run_font(r, size=11, color=BLUE, bold=True)
    r = p.add_run("  娇兰 KOL PDF Brief → 第三代黄金复原蜜品牌片")
    set_run_font(r, size=11, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("图文广告主案例")
    set_run_font(r, size=11, color=BLUE, bold=True)
    r = p.add_run("  策略 Brief → 小红书 3:4 三图素材")
    set_run_font(r, size=11, color=INK)

    add_callout(
        doc,
        "本轮范围说明",
        "语音/配音不是本轮演示门禁。文本、图片、视频、视觉理解、联网研究和文档解析可用即可开始；语音不可用时允许 Fixture、静音或跳过音频步骤，但不得阻断视频主链路。",
        fill=AMBER,
        label_color=AMBER_TEXT,
    )
    doc.add_page_break()


def add_overview(doc):
    doc.add_heading("1. 一分钟看懂：测什么、怎么走、何时停止", level=1)
    add_label_paragraph(
        doc,
        "测试目标",
        "验证同一 Project 下，从 Brief、策略审批、跨系统 Handoff 到 CreativeTask、资产生成、质量检查和交付包的血缘完整、状态可恢复、结果可审计。",
    )

    add_table(
        doc,
        ["链路", "起点", "中间必须产物", "终点"],
        [
            ["品牌广告", "娇兰 PDF Brief", "已批准 StrategyPackage → brand_video Handoff → Brand Brief → Direction → BrandFilm", "已交付品牌视频 CreativePackage"],
            ["图文广告", "独立小红书 Brief", "已批准 StrategyPackage → image_text Handoff → Direction → Draft → 3 个成品槽位", "已交付 3:4 图文 CreativePackage"],
        ],
        [1350, 1980, 3930, 2100],
        font_size=9.1,
        center_columns=(0,),
    )

    doc.add_heading("1.1 执行规则", level=2)
    add_list(
        doc,
        [
            "每个阶段只做本节列出的动作，不跨步骤创建下游任务。",
            "生成按钮只点击一次；等待状态从提交/排队变化后再操作。刷新不是重试。",
            "每阶段至少记录一个可追溯 ID 或 Hash，并保存关键页面截图。",
            "若出现跨 Project、数据丢失、错误 Handoff 或不可恢复白屏，立即停止并按 P0 反馈。",
            "若上游阶段失败，下游阶段标记“阻塞”，不要用手工复制粘贴绕过。",
        ],
        kind="decimal",
    )

    doc.add_heading("1.2 结果怎么判", level=2)
    add_table(
        doc,
        ["结论", "使用条件", "下一步"],
        [
            ["通过", "所有“应该看到”均满足，且关键证据已记录", "进入下一阶段"],
            ["失败", "实际结果与预期不一致，可稳定复现", "保存证据并反馈"],
            ["阻塞", "上游、权限、数据或 Provider 使本阶段无法开始", "不要绕行；反馈阻塞点"],
            ["不适用", "明确不在本轮范围，例如语音真实合成", "写明原因后继续"],
        ],
        [1200, 5160, 3000],
        font_size=9.5,
        center_columns=(0,),
    )

    doc.add_heading("1.3 问题反馈最小信息", level=2)
    add_callout(
        doc,
        "遇到问题就发这些",
        "阶段编号 + 当前 URL + 操作前状态 + 实际提示/错误码 + Request ID/任务 ID + 一张截图。若刷新或重试过，还要说明刷新/重试后的变化。",
        fill=RED,
        label_color=RED_TEXT,
    )


def add_preflight(doc):
    doc.add_heading("2. 前置检查：开始测试前先过四道门", level=1)
    add_callout(
        doc,
        "开始条件",
        "健康、登录权限、Migration、在本轮范围内的 Provider 全部通过后开始。语音不在本轮范围，不影响 Provider 总结论。",
        fill=GREEN,
        label_color=GREEN_TEXT,
    )
    rows = [
        ["PF-01", "前端与健康", "打开测试 URL；访问 /healthz、/readyz", "首页可访问；两端点 HTTP 200；状态为 ok / ready", "通过"],
        ["PF-02", "登录与权限", "登录后进入目标 Project，读取策略任务与创意任务", "组织/用户/成员 active；具备 project.read/write、creative.read/write", "通过"],
        ["PF-03", "Project / Brand", "检查 Project 状态与 Brand 绑定", "Project active；brand_id 非空；上下文与页面一致", "通过"],
        ["PF-04", "Migration", "核对部署流程与最新领域接口", "部署启动前强制运行 cookies-migrate；Strategy/Creative/Project 接口可读", "通过"],
        ["PF-05", "Provider", "读取 Provider 与 Strategy capabilities", "文本、图片、视频、视觉理解、联网研究、文档视觉 available=true", "通过"],
        ["PF-06", "语音", "读取 speech capability", "本轮允许 available=false；必须有 Fixture/跳过说明且不阻断视频", "不适用"],
    ]
    table = add_table(
        doc,
        ["编号", "测试什么", "怎么做", "应该看到", "当前基线"],
        rows,
        [900, 1450, 2300, 3610, 1100],
        font_size=8.8,
        center_columns=(0, 4),
    )
    for row in table.rows[1:]:
        status = row.cells[4].text
        shade_cell(row.cells[4], GREEN if status == "通过" else AMBER)
        style_cell_text(row.cells[4], color=GREEN_TEXT if status == "通过" else AMBER_TEXT,
                        size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("2.1 本轮已核验环境", level=2)
    add_table(
        doc,
        ["项目", "证据"],
        [
            ["环境", "14.103.24.58:8091（HTTP 测试环境）；/healthz=ok；/readyz=ready"],
            ["身份", "org_local / user_local；owner；组织、用户、成员均 active"],
            ["Project", "project_local；active；已绑定 Local Brand；Context Version=1"],
            ["权限", "project.read/write、creative.read/write 等服务端 Scope 已授予"],
            ["Provider", "文本、图片、视频、视觉理解、联网研究、文档视觉均 configured + available"],
            ["语音", "当前为 Fixture、available=false；不在本轮验收范围"],
        ],
        [1800, 7560],
        font_size=9.5,
    )


def add_flow_maps(doc):
    doc.add_page_break()
    doc.add_heading("3. 两条主链路总览", level=1)
    add_callout(doc, "门禁规则", "上一阶段的“阶段输出”是下一阶段的唯一可信输入。缺少上游 ID、Version 或 Hash 时，不得手工补造下游对象。")

    doc.add_heading("3.1 品牌广告链路", level=2)
    add_table(
        doc,
        ["阶段", "测试什么", "阶段输出"],
        [
            ["BA-01", "Project、Brand、Product 上下文与隔离", "正确 Project 上下文"],
            ["BA-02", "PDF 上传、解析、持久化", "KnowledgeDocument / 解析结果"],
            ["BA-03", "区分事实与缺失项", "持久化对话与 Brief 事实"],
            ["BA-04", "联网证据不覆盖 Brief", "带来源的补充证据"],
            ["BA-05", "确认单品 Brief，生成并批准策略", "已批准 StrategyPackage + Hash"],
            ["BA-06", "brand_video Route 与跨系统 Handoff", "Handoff + Intake + Route"],
            ["BA-07", "Brand Brief 与品牌方向", "Confirmed Brand Brief + Direction"],
            ["BA-08", "渠道任务、方案、脚本与分镜", "BrandFilm Plan + Units"],
            ["BA-09", "单元生成、锁定与预览恢复", "Preview Asset + 完整 Attempts"],
            ["BA-10", "质量、批准与交付", "CreativeVersion + CreativePackage"],
        ],
        [1100, 4750, 3510],
        font_size=9.1,
        center_columns=(0,),
    )

    doc.add_heading("3.2 图文广告链路", level=2)
    add_table(
        doc,
        ["阶段", "测试什么", "阶段输出"],
        [
            ["IT-01", "独立 Project 与结构化 Brief", "Confirmed Brief"],
            ["IT-02", "图文策略生成、评审、批准", "已批准 StrategyPackage + Hash"],
            ["IT-03", "image_text Route 与 Handoff", "Handoff + Intake"],
            ["IT-04", "3 个方向与任务创建", "Selected Direction + CreativeTask"],
            ["IT-05", "Draft、文案与 3 个槽位", "可生成 Draft Revision"],
            ["IT-06", "逐槽位底图和中文排版", "3 个成功 AssetVersionRef"],
            ["IT-07", "单槽重生成、采用与原子物化", "Final Draft Revision"],
            ["IT-08", "冻结、检查、批准、交付和恢复", "CreativePackage + 完整 Lineage"],
        ],
        [1100, 4750, 3510],
        font_size=9.1,
        center_columns=(0,),
    )


def add_stage(doc, stage_id, title, priority, test_goal, prerequisite, steps, expectations,
              output, evidence, note=None):
    doc.add_heading(f"{stage_id}  {title}", level=2)
    add_label_paragraph(doc, "测试什么", test_goal)
    add_callout(doc, "开始前", prerequisite, fill=LIGHT_GRAY, label_color=MUTED)
    if note:
        add_callout(doc, "本阶段说明", note, fill=AMBER, label_color=AMBER_TEXT)
    doc.add_heading("怎么做", level=3)
    add_list(doc, steps, kind="decimal", after=3)
    doc.add_heading("结果应该是什么", level=3)
    rows = []
    for index, expected in enumerate(expectations, start=1):
        rows.append([str(index), expected, "□ 通过  □ 失败\n证据：________________"])
    add_table(doc, ["#", "应该看到", "执行记录"], rows, [650, 6140, 2570],
              font_size=9.0, center_columns=(0,))
    add_label_paragraph(doc, "阶段输出", output, label_color=GREEN_TEXT, after=3)
    add_label_paragraph(doc, "关键证据", evidence, label_color=DARK_BLUE, after=3)
    add_label_paragraph(doc, "阶段结论", f"□ 通过  □ 失败  □ 阻塞  □ 不适用     优先级：{priority}",
                        label_color=BLUE, after=10)


def add_brand_chain(doc):
    doc.add_page_break()
    doc.add_heading("4. 主链路 A：策略对接品牌广告", level=1)
    add_callout(
        doc,
        "主案例",
        "娇兰 KOL 总 Brief → 确认主推第三代黄金复原蜜 → brand_video → 抖音 30 秒 9:16 品牌种草视频。",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("4.1 固定测试数据", level=2)
    add_table(
        doc,
        ["字段", "固定值"],
        [
            ["Project", "优先使用专用娇兰验收 Project；若使用 project_local，先确认 Brand/Product 绑定正确"],
            ["Brand / Product", "法国娇兰 / 娇兰第三代黄金复原蜜"],
            ["Brief 文件", "【娇兰】brief(1).pdf；多产品、双十一种草 KOL 总 Brief"],
            ["目标人群", "25-35 岁，关注抗老修护与干皮护理的高端护肤人群"],
            ["核心主张", "稳筑「肌」本功，愈现年轻态"],
            ["目标 Route", "brand_video；douyin；30 秒；9:16"],
            ["视觉调性", "高端、温润、克制、可信；金色暖调、明亮简洁、产品真实"],
            ["禁止项", "不得混入其他产品卖点；不得称为油；不得使用虚假 before/after 或医学绝对化承诺"],
        ],
        [1900, 7460],
        font_size=9.3,
    )

    add_stage(
        doc, "BA-01", "Project 与品牌上下文", "P0",
        "验证进入的 Project、Brand、Product 正确，并且刷新、直接 URL 与跨页面访问不会串 Project。",
        "已登录；目标 Project 已创建、状态 active、已绑定 Brand。",
        [
            "进入目标 Project 的“需求与策略 → 策略工作区”。",
            "核对页面顶部的 Project、品牌和产品信息。",
            "刷新页面；复制当前 URL，在新标签页直接打开。",
            "观察页面是否出现其他 Project 的 Brief、消息、任务或资产。",
        ],
        [
            "Project、Brand、Product 与本次测试数据一致。",
            "页面只显示当前 Project 的业务对象，不出现其他项目内容。",
            "刷新和直接访问 URL 后仍恢复同一 Project，加载状态能结束。",
            "不出现 project must be active and brand-bound、永久骨架屏或无权限误报。",
        ],
        "正确的 ProjectContext，可作为 PDF、对话、策略和 Creative 的共同作用域。",
        "Project ID、Brand ID、Product ID、URL、刷新前后截图。",
    )

    add_stage(
        doc, "BA-02", "上传并解析娇兰 PDF Brief", "P0/P1",
        "验证 PDF 可上传、解析、检索和持久化，并且解析不会阻塞对话页面。",
        "BA-01 通过；准备【娇兰】brief(1).pdf。",
        [
            "点击“添加资料”，上传【娇兰】brief(1).pdf。",
            "等待状态从上传/解析中变为可用，期间确认输入框仍可操作。",
            "打开“来源资料”，核对文件名、解析状态、片段数量或预览。",
            "刷新页面后再次打开来源资料。",
        ],
        [
            "上传后出现明确解析状态，页面不会冻结。",
            "来源资料显示文件可用，文件名正确，刷新后仍存在。",
            "解析结果识别这是多产品 KOL 总 Brief，不擅自确定主推产品。",
            "第三代黄金复原蜜、蜂皇水、复原霜、粉底液的事实边界可区分。",
            "医学或绝对化高风险表述进入待复核，不成为广告承诺。",
        ],
        "可被对话和策略引用的 KnowledgeDocument。",
        "KnowledgeDocument ID、文件 SHA256、解析状态、片段数、页面截图。",
    )

    add_stage(
        doc, "BA-03", "普通对话识别事实与缺失项", "P1",
        "验证助手能基于 PDF 分开“已知事实”和“仍需确认”，且不会串用多产品卖点。",
        "BA-02 通过；关闭联网搜索。",
        [
            "发送以下输入：请根据我上传的 Brief，总结品牌、沟通主题、候选产品、发布平台、发布形式和拍摄限制。请把 PDF 已明确的事实与仍然缺失的信息分开，暂时不要替我选择主推产品。",
            "检查回复中的品牌、主题、候选产品、平台、形式、限制和待确认项。",
            "继续发送一条追问后刷新页面。",
        ],
        [
            "回复明确这是多产品 KOL 总 Brief。",
            "事实与待确认项分开；KOL、主推产品、发布时间、时长等未被编造。",
            "四个产品卖点未串用，且仅引用当前 PDF。",
            "消息和回复落库，刷新后仍存在，可继续对话。",
        ],
        "持久化的对话事实和待确认问题，为 Brief 确认提供输入。",
        "Conversation ID、Message ID、回复截图、刷新后截图。",
    )

    add_stage(
        doc, "BA-04", "联网核验且保持来源边界", "P1",
        "验证联网研究只补充有来源支持的信息，不覆盖 PDF，也不冒充用户确认。",
        "BA-03 通过；联网研究能力 available=true。",
        [
            "开启“联网搜索 / 搜索后回答”。",
            "发送以下输入：请联网核验娇兰第三代黄金复原蜜的公开品牌信息和官方产品表述，只补充有可靠来源支持的内容。不要把联网信息覆盖为 Brief 原文，也不要补写 PDF 中尚未确定的 KOL、发布时间或渠道支持。",
            "观察排队/搜索状态；完成后查看来源和最终回复。",
            "进行中和完成后各刷新一次。",
        ],
        [
            "用户消息先落库，搜索完成前不会生成无来源的最终回复。",
            "完成后只出现一条基于来源的回复，结论与来源可对应。",
            "优先使用官方或可信来源；联网信息与 Brief 事实有明确边界。",
            "刷新不会重复创建搜索任务；失败时可重试且不阻断普通对话。",
        ],
        "带来源的补充证据，可供策略生成引用。",
        "Research/Search Run ID、来源 URL、Message ID、Request ID。",
    )

    add_stage(
        doc, "BA-05", "确认单品 Brief 并生成已批准策略", "P0/P1",
        "验证用户确认能收敛到第三代黄金复原蜜，并生成可追溯、可评审、不可变的策略包。",
        "BA-02 至 BA-04 通过。",
        [
            "发送确认文本：本轮主推“娇兰第三代黄金复原蜜”，目标受众为 25-35 岁、关注抗老修护和干皮护理的人群，发布平台为抖音，内容形式为 30 秒品牌种草视频，目标是双十一期间建立产品认知并促进种草。核心主张为“稳筑「肌」本功，愈现年轻态”。其余未确认信息继续保留为待确认。",
            "进入 Brief 页面，核对品牌、产品、受众、目标、渠道、限制和证据后确认 Brief。",
            "进入策略页生成策略，完成评审并批准版本。",
            "刷新页面，重新打开已批准版本。",
        ],
        [
            "Brief 与当前 Brand/Product 兼容，且只围绕第三代黄金复原蜜。",
            "Mandatory、Prohibited、Claims、Evidence、Open Questions 均有内容。",
            "实验结论保留样本和限制，高风险说法不进入对外承诺。",
            "生成不可变 StrategyPackageVersion、Content Hash 和 approved 状态。",
            "刷新后仍是同一已批准版本，内容和 Hash 不变。",
        ],
        "Approved StrategyPackageVersion，是 Handoff 的唯一合法上游。",
        "Brief ID/Version、Strategy ID/Version、Package ID/Version/Hash、审批截图。",
    )

    add_stage(
        doc, "BA-06", "选择 brand_video Route 并执行 Handoff", "P0",
        "验证跨系统交接携带冻结的策略引用、Route、渠道规格和完整 Hash，不靠复制粘贴。",
        "BA-05 通过；Route readiness=ready。",
        [
            "进入“创意任务策略 / 创意交接”。",
            "选择品牌广告 / brand_video，渠道选择 douyin，核对 30 秒、9:16。",
            "点击交接一次，等待进入品牌广告工作区。",
            "刷新页面并再次打开同一任务；不要重复点击创建。",
        ],
        [
            "Route 显示 brand_video、douyin、30 秒、9:16，并给出选择理由。",
            "Handoff 携带 Package Ref/Hash、Contract Version、Handoff Hash。",
            "上游 readiness=ready；阻塞项为空或有明确修复入口。",
            "刷新后恢复同一 Intake/Task；重复请求幂等，不创建重复任务。",
        ],
        "Creative Handoff、Intake 和 brand_video Route。",
        "Route ID、Handoff Hash、Intake ID、Idempotency Key、页面 URL。",
    )

    add_stage(
        doc, "BA-07", "确认 Brand Brief 并选择品牌方向", "P0/P1",
        "验证 Creative 能继承策略与资产边界，Brand Brief 可编辑、可确认，候选方向真正有差异。",
        "BA-06 通过；品牌 Intake 已创建；文本模型可用。",
        [
            "进入 Brand Brief Review，核对品牌、产品、场景、卖点、证据、资产候选和限制。",
            "修正缺失项；确认商品图/Logo 授权；保存并确认 Brief。",
            "点击“生成品牌方向”一次；完成后比较全部候选。",
            "选择一个方向并确认；刷新页面。",
        ],
        [
            "Brand Brief 的品牌和产品正确，用户编辑不会被刷新或后端投影覆盖。",
            "未确认资产或权利时有明确 blocker；清除后才可确认并生成 Brief Hash。",
            "至少 2-3 个方向在概念、叙事、情绪弧、视觉语法上真正不同。",
            "所有方向继承核心主张、Mandatory/Prohibited 与证据边界。",
            "确认后生成不可变 Direction ID/Hash，刷新后仍选中同一方向。",
        ],
        "Confirmed Brand Brief + Selected Direction。",
        "Brand Brief Revision/Hash、Direction Batch、Direction ID/Hash、资产引用。",
    )

    add_stage(
        doc, "BA-08", "创建渠道任务并生成方案、脚本与分镜", "P1",
        "验证渠道任务继承全部上游血缘，并生成满足时长、画幅、品牌和合规约束的 BrandFilm Plan。",
        "BA-07 通过；douyin Route ready。",
        [
            "用已确认 Direction 创建抖音 Brand Video Task；只创建一次。",
            "生成品牌片方案，检查脚本、分镜、时长、场景和产品使用方式。",
            "确认方案并准备生成单元。",
            "如页面显示音频步骤，可查看但不要求真实语音合成。",
        ],
        [
            "Task Source Snapshot 记录 Strategy Package、Handoff、Brand Brief、Direction Hash。",
            "脚本体现核心主张、卖点、使用方式和禁用项，不混入其他产品。",
            "分镜总时长约 30 秒、9:16，镜头与脚本逐段对应。",
            "商品外观、Logo、包装真实和金色暖调进入镜头约束。",
            "语音不可用时显示 Fixture/跳过提示，不阻断 Plan、Units 或视频生成。",
        ],
        "Brand Video Task + BrandFilm Plan + Generation Units。",
        "CreativeTask ID/Version、Source Snapshot、Plan Hash、Unit IDs。",
        note="语音不是本轮阻断项；若语音不可用，本阶段仍可判通过，但应记录为“不适用/Fixture”。",
    )

    add_stage(
        doc, "BA-09", "生成单元、锁定并合成预览", "P1",
        "验证每个生成单元独立、可重试、可选择 Attempt，刷新或重启后不丢任务和血缘。",
        "BA-08 通过；参考资产已授权；图片/视频 Provider 可用。",
        [
            "逐单元提交生成；每次只点击一次，观察 ProviderJob 状态。",
            "若单元失败，仅对该单元重试；选择满意 Attempt 并锁定。",
            "所有单元锁定后合成预览。",
            "在生成中刷新一次；如环境允许，再执行一次 API 重启恢复检查。",
        ],
        [
            "每个 Unit 有独立 Attempt、ProviderJob 和 AssetVersionRef。",
            "单个失败不污染其他单元；重试只新增目标单元 Attempt。",
            "锁定只接受当前 Revision 和当前 Attempt。",
            "预览为 9:16、约 30 秒、可播放，与选定 Direction 一致。",
            "刷新/重启后任务继续或明确失败并可恢复，Lineage 不丢失。",
        ],
        "可供质量检查的 BrandFilm Preview Asset。",
        "Unit/Attempt/Job/Asset IDs、Provider、提交/完成时间、预览 Asset。",
    )

    add_stage(
        doc, "BA-10", "质量检查、人工确认与交付", "P0/P1",
        "验证质量门禁、人工确认、版本冻结、批准和交付顺序正确，最终包不可变且血缘完整。",
        "BA-09 通过；品牌片预览已完成。",
        [
            "运行自动质量检查。",
            "人工检查品牌一致、产品真实、功效合规、画面完整和渠道规格。",
            "确认质量；依次 Finalize、Approve、Deliver。",
            "刷新页面并重新打开最终交付包。",
        ],
        [
            "自动检查结果可解释，失败有证据和修复建议。",
            "必需人工检查未完成时不得 Finalize/Approve。",
            "状态按质量确认 → finalized → approved → delivered 前进。",
            "最终 CreativePackage 不可变，包含视频与完整 Lineage。",
            "交付包仍引用同一 Strategy Package、Brand Brief 和 Direction Hash。",
        ],
        "Delivered Brand Video CreativePackage。",
        "Quality Run、Manual Checks、CreativeVersion ID、CreativePackage ID、Lineage。",
    )


def add_image_text_chain(doc):
    doc.add_page_break()
    doc.add_heading("5. 主链路 B：策略对接图文广告", level=1)
    add_callout(
        doc,
        "隔离原则",
        "图文链路使用独立 Project 和独立 Brief，不复用娇兰品牌视频任务。若使用同一品牌，也必须新建独立 Strategy Route 与 CreativeTask。",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("5.1 固定测试 Brief", level=2)
    add_table(
        doc,
        ["字段", "固定值"],
        [
            ["Project / Brand", "青柠气泡水小红书图文全链路验收 / 清醒气泡（测试品牌）"],
            ["Product", "0 糖青柠气泡水"],
            ["目标 / 受众", "建立新品认知，承接搜索与收藏 / 20-32 岁城市通勤人群"],
            ["渠道 / 规格", "xiaohongshu / image_text / 3:4 / 1080×1440 PNG"],
            ["CTA", "搜索“清醒气泡”，了解更多新品信息"],
            ["Mandatory", "品牌名、授权商品正面图、0 糖事实、商业内容披露"],
            ["Prohibited", "不得虚构减脂/治疗、检测、销量、用户评价或促销"],
            ["图片槽位", "01 封面 / 02 证据 / 03 CTA；三个独立 AssetVersion"],
        ],
        [1900, 7460],
        font_size=9.3,
    )
    add_callout(
        doc,
        "可直接复制的策略输入",
        "请为测试品牌“清醒气泡”的“0 糖青柠气泡水”生成一份小红书图文策略。目标受众是 20-32 岁城市通勤人群，传播目标是建立新品认知并承接品牌词、品类词搜索与收藏。核心信息是“0 糖青柠气泡水适合通勤、午后和运动后清爽饮用”。内容必须包含品牌名、授权商品正面图、0 糖事实和商业内容披露；不得虚构减脂、减肥、治疗功效、检测数据、销量、用户评价或促销。输出 Route 必须为 xiaohongshu / image_text / 3:4 / 1080×1440，CTA 为“搜索清醒气泡，了解更多新品信息”。",
        fill=CALLOUT,
    )

    add_stage(
        doc, "IT-01", "创建独立 Project 与结构化 Brief", "P0",
        "验证图文链路从独立 Project/Brand/Brief 开始，结构化字段完整，且不会出现娇兰数据。",
        "已登录；Project 服务可用。",
        [
            "新建独立 Project，并填写上方 Project、Brand、Product 和目标。",
            "进入策略工作区，粘贴固定策略输入。",
            "检查结构化 Brief 的受众、目标、渠道、规格、CTA、Mandatory 和 Prohibited。",
            "确认 Brief 后刷新页面。",
        ],
        [
            "Project 为 active 且 brand-bound，可正常进入。",
            "Brief 字段完整，Route 意图为 xiaohongshu/image_text/3:4。",
            "不出现娇兰 Project、Brief、策略、消息或资产。",
            "刷新后 Brief、Version 和确认状态仍存在。",
        ],
        "Confirmed Image Text Brief。",
        "Project/Brand/Product ID、Brief ID/Version、确认截图。",
    )

    add_stage(
        doc, "IT-02", "生成、评审并批准图文策略", "P0/P1",
        "验证策略覆盖人群、场景、信息顺序、搜索意图、互动意图、限制和固定图文规格。",
        "IT-01 通过。",
        [
            "基于已确认 Brief 生成完整策略。",
            "检查内容角度、信息优先级、搜索/收藏意图、CTA 和图片槽位职责。",
            "完成评审并批准 StrategyPackageVersion。",
            "刷新并重新打开已批准版本。",
        ],
        [
            "策略包含目标人群、生活场景、核心信息与证据、CTA 意图。",
            "Route 为 xiaohongshu_image_text / image_text / 3:4。",
            "禁止项和商业内容披露进入 Production Requirements。",
            "已批准版本带 Package ID、Version、Content Hash，刷新后不变。",
        ],
        "Approved Image Text StrategyPackageVersion。",
        "Strategy/Package ID、Version、Hash、审批记录。",
    )

    add_stage(
        doc, "IT-03", "验证 image_text Handoff", "P0",
        "验证 Handoff 只使用当前 Project 的冻结策略引用，并完整保留 Route、规格、授权和限制。",
        "IT-02 通过；Route readiness=ready。",
        [
            "进入创意任务策略，选择“小红书图文”。",
            "核对 xiaohongshu / image_text / 3:4 / 1080×1440。",
            "点击 Handoff 一次，进入图文创作。",
            "刷新页面并查看 Strategy → Creative 交接摘要。",
        ],
        [
            "Handoff 只使用当前 Project 的冻结 Refs。",
            "显示 Package Ref、Handoff Hash、Route ID、渠道和规格。",
            "商品图授权、Mandatory、Prohibited 不丢失。",
            "重复请求幂等，不产生重复 Intake。",
        ],
        "Image Text Handoff + Creative Intake。",
        "Handoff/Route/Intake ID、Hash、Idempotency Key。",
    )

    add_stage(
        doc, "IT-04", "生成并选择 3 个 Creative Direction", "P1",
        "验证候选方向在概念和执行上真实不同，并且确认后创建同一血缘下的 CreativeTask。",
        "IT-03 通过；图文 Intake ready；文本模型可用。",
        [
            "点击“生成 3 个候选方向”一次。",
            "对比概念、理由、执行提纲、视觉关键词、CTA 和边界。",
            "选择一个方向并确认创建图文任务。",
            "刷新并直接访问该任务 URL。",
        ],
        [
            "准确生成 3 个可区分方向，而不是只更换标题。",
            "每个方向均绑定 Strategy Handoff 与 Route。",
            "方向包含内容角度、信息顺序、视觉关键词、CTA 和限制。",
            "确认后 URL 与刷新都恢复同一 CreativeTask。",
        ],
        "Selected Direction + Image Text CreativeTask。",
        "Direction Batch、Direction ID/Hash、CreativeTask ID。",
    )

    add_stage(
        doc, "IT-05", "生成图文方案并编辑文案", "P1",
        "验证 Draft 结构、文案编辑、版本校验和三个槽位职责正确，刷新后不丢修改。",
        "IT-04 通过；CreativeTask 已创建。",
        [
            "点击“生成图文方案”一次。",
            "检查 3 个标题、正文、话题、封面文案和三个图片槽位。",
            "修改选中标题、正文和每个槽位 Overlay Copy，保存。",
            "刷新页面并再次核对。",
        ],
        [
            "Draft 为 creative-image-text-draft/v2，包含 Direction Ref 与 Input Identity Hash。",
            "三个槽位固定为封面/证据/CTA，顺序和职责明确。",
            "保存校验 Expected Task Version 和 Expected Draft Revision。",
            "刷新后修改仍存在；开始图片生成后输入冻结或保存按钮禁用。",
        ],
        "可生成的 Image Text Draft Revision。",
        "Task Version、Draft Revision、Direction Ref、保存前后截图。",
    )

    add_stage(
        doc, "IT-06", "逐槽位生成底图与排版成品", "P1",
        "验证三个槽位独立生成，模型底图与中文排版职责分离，成品规格和资产引用稳定。",
        "IT-05 通过；商品图已授权；图片 Provider 与中文字体可用。",
        [
            "按槽位 1、2、3 依次点击“生成这一张”；每槽只点击一次。",
            "观察等待 → 底图生成 → 排版合成 → 成品就绪。",
            "每个槽位完成后打开成品预览。",
            "记录每槽的 Attempt、ProviderJob、底图和成品 Asset。",
        ],
        [
            "每次请求绑定 Task Version、Draft Revision 和独立 Idempotency Key。",
            "模型底图无字；中文由服务端模板排版，无乱码或豆腐块。",
            "三个成品均为 1080×1440 PNG，并有稳定 AssetVersionRef。",
            "三个槽位互不覆盖；单槽失败不影响其他槽位。",
            "首次成功 Attempt 自动采用。",
        ],
        "3 个槽位的成功 AssetVersionRef。",
        "每槽 PromptPackage/Attempt/ProviderJob/底图 Asset/成品 Asset。",
    )

    add_stage(
        doc, "IT-07", "单槽重生成、人工采用与原子物化", "P0/P1",
        "验证重生成只影响目标槽位，采用有版本保护；三槽齐备后一次性物化到新 Draft Revision。",
        "IT-06 通过；至少一个槽位已有采用结果，最终三个槽位均成功。",
        [
            "选择一个槽位，点击“重新生成这一张”。",
            "等待新 Attempt 成功并采用；刷新页面。",
            "检查旧 Attempt 仍可追溯，其他槽位没有新增 Attempt。",
            "三槽均采用后刷新 Workspace，核对 Final Draft Revision。",
        ],
        [
            "只有目标槽位新增 Attempt，其他槽位不重建。",
            "采用请求带 Expected Task Version 与 Expected Selection Version。",
            "旧 Attempt 可追溯，新采用结果刷新后保持；旧页面冲突不能覆盖新选择。",
            "任务进入 ready_for_review，三个成品一次性物化到一个新 Draft Revision。",
            "Draft、Slots、Assets 与 Direction/Strategy Lineage 一致。",
        ],
        "Final Image Text Draft Revision，包含 3 个成品引用。",
        "Slot Order、Attempt IDs、Selection Version、Final Draft Revision、3 个 AssetVersionRef。",
    )

    add_stage(
        doc, "IT-08", "冻结、检查、批准、交付与恢复", "P0/P1",
        "验证版本状态机、交付门禁、最终包和刷新/直接 URL/重启恢复能力。",
        "IT-07 通过。",
        [
            "冻结当前版本并执行交付检查。",
            "确认检查通过后人工批准，并生成交付包。",
            "生成中和交付后各刷新一次，并直接访问任务 URL。",
            "如环境允许，在生成或排版过程中执行一次 API 重启恢复检查。",
        ],
        [
            "状态依次为 created → checked → approved → delivered。",
            "未通过检查时不得批准；批准后版本不可变。",
            "最终 CreativePackage 包含三张成品和发布元数据。",
            "交付包保留 Strategy Package、Handoff、Direction、Draft、Asset Lineage。",
            "刷新/直接 URL/重启不会重复提交 Job 或丢失 Attempt，交付包仍可读取。",
        ],
        "Delivered Image Text CreativePackage。",
        "CreativeVersion ID、Check Run、Approval、CreativePackage ID、恢复前后截图。",
    )


def add_negative_tests(doc):
    doc.add_page_break()
    doc.add_heading("6. 必测异常与安全边界", level=1)
    add_callout(
        doc,
        "执行原则",
        "异常用例在主链路至少跑通一次后执行。不要为制造异常而修改真实客户数据；优先使用测试 Project、旧标签页、重复请求和受控失败场景。",
        fill=AMBER,
        label_color=AMBER_TEXT,
    )
    add_table(
        doc,
        ["编号", "测试什么", "怎么做", "结果应该是什么", "结果"],
        [
            ["NEG-01", "版本冲突", "A 页面保存后，用旧 B 页面保存或确认", "B 返回冲突，不覆盖 A 的新版本", "待测"],
            ["NEG-02", "重复请求", "重复点击或重放相同 Idempotency-Key", "相同输入返回原资源；不产生重复任务/扣费", "待测"],
            ["NEG-03", "单 Job 失败", "让一个图片/视频 Attempt 失败", "仅目标 Attempt/Unit 失败；其他对象不受影响，可重试", "待测"],
            ["NEG-04", "刷新/直接 URL", "在排队、生成、完成状态刷新并复制 URL 打开", "恢复同一任务，不重复提交，不永久骨架屏", "待测"],
            ["NEG-05", "跨 Project", "用 Project A 的资源 ID 在 Project B 路由访问", "拒绝访问，不泄露对象存在性", "待测"],
            ["NEG-06", "缺失必填", "清空 Brief/Route 的必填字段", "显示 blocker；不能生成、批准或交付", "待测"],
            ["NEG-07", "授权阻断", "使用未授权商品图或 generative_ai_allowed=false", "生成按钮禁用，不绕过授权", "待测"],
            ["NEG-08", "非法路由", "访问不存在的 task/direction/route URL", "返回受控 404/错误页，不白屏", "待测"],
        ],
        [900, 1650, 2380, 3380, 1050],
        font_size=8.4,
        center_columns=(0, 4),
    )


def add_evidence_and_exit(doc):
    doc.add_heading("7. 全链路证据台账", level=1)
    add_table(
        doc,
        ["对象", "品牌广告记录", "图文广告记录"],
        [
            ["Project / Brand / Product", "________________", "________________"],
            ["Brief ID / Version", "________________", "________________"],
            ["Strategy / Package / Hash", "________________", "________________"],
            ["Handoff / Route / Hash", "________________", "________________"],
            ["Intake / Direction", "________________", "________________"],
            ["CreativeTask / Version", "________________", "________________"],
            ["Draft / BrandFilm Revision", "________________", "________________"],
            ["Attempt / ProviderJob", "________________", "________________"],
            ["AssetVersionRef", "________________", "________________"],
            ["Quality / Approval", "________________", "________________"],
            ["CreativeVersion / Package", "________________", "________________"],
        ],
        [2500, 3430, 3430],
        font_size=9.0,
    )

    doc.add_heading("8. 缺陷反馈模板", level=1)
    add_table(
        doc,
        ["字段", "填写内容"],
        [
            ["标题", "[P0/P1/P2/P3] 阶段编号 - 问题简述"],
            ["环境", "前端 URL / API / Commit / Provider / 浏览器"],
            ["业务对象", "Project ID；Brief/Package/Handoff/Task/Job ID"],
            ["复现步骤", "1. ______  2. ______  3. ______"],
            ["预期结果", "文档本阶段“结果应该是什么”中的对应条目"],
            ["实际结果", "页面实际状态、提示或错误"],
            ["错误证据", "HTTP Status / Error Code / Request ID / 截图"],
            ["复现情况", "□ 必现  □ 偶现；次数：____/____"],
            ["刷新/重试", "刷新后：________  重试后：________"],
        ],
        [2100, 7260],
        font_size=9.3,
    )

    doc.add_heading("9. 最终结论与退出标准", level=1)
    add_table(
        doc,
        ["结论项", "填写"],
        [
            ["品牌广告", "□ 通过  □ 有条件通过  □ 不通过"],
            ["图文广告", "□ 通过  □ 有条件通过  □ 不通过"],
            ["整体结论", "□ 通过  □ 有条件通过  □ 不通过"],
            ["P0/P1 缺陷", "____________________________________________"],
            ["遗留风险", "____________________________________________"],
            ["签字 / 日期", "测试：______  产品：______  研发：______  日期：______"],
        ],
        [2100, 7260],
        font_size=9.5,
    )
    add_label_paragraph(doc, "退出标准", "以下条件全部满足才能判定全链路通过。")
    add_list(
        doc,
        [
            "两条链路均从当前 Project 的 Confirmed Brief 开始，并生成 Approved StrategyPackage。",
            "品牌广告进入 brand_video；图文广告进入 image_text；Route 和渠道规格正确。",
            "Package、Handoff、Direction、Task、Revision、Asset、CreativePackage 血缘完整且 Hash 一致。",
            "刷新、直接 URL、重复请求和版本冲突不会造成重复任务、数据丢失或串 Project。",
            "品牌广告完成预览、质量、批准和交付；图文广告完成三槽、原子物化、冻结、批准和交付。",
            "不存在 P0；不存在阻断主链路的 P1；遗留问题均记录负责人和处理计划。",
        ],
        kind="bullet",
    )


def audit_document(doc):
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert len(doc.tables) > 10
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa"
        assert int(tbl_w.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA
        for row in table.rows:
            assert len(row.cells) == len(grid_widths)
            for idx, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tc_w is not None and int(tc_w.get(qn("w:w"))) == grid_widths[idx]


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_overview(doc)
    add_preflight(doc)
    add_flow_maps(doc)
    add_brand_chain(doc)
    add_image_text_chain(doc)
    add_negative_tests(doc)
    add_evidence_and_exit(doc)
    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
